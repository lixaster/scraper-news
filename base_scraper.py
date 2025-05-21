import os
import re
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import requests
from utils_func import setup_logging, load_config


class BaseScraper:
    # 包括静态属性和方法
    def __init__(self):
        # 配置日志
        self.logger = setup_logging()
        # 读取配置文件
        config = load_config()
        self.config = config
        # 获取配置文件中的新闻站点
        for news_site in config.get("news_sites"):
            if news_site.get("name") == self.source_name:
                self.source_name_cn = news_site.get("name_cn")
                self.url = news_site.get("url")
                break
        self.logger.info(f"{self.source_name_cn} 爬虫开始运行...")

        self.uid = int(config.get("uid", 1027))
        self.gid = int(config.get("gid", 100))
        self.notify_switch = config.get("notify_switch", False)
        self.webhook_url = config.get("webhook_url", None)
        self.save_folder = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            config.get("save_folder", "docx"),
            self.source_name,
        )
        self.playwright_timeout = int(config.get("playwright_timeout", 2)) * 60000
        os.makedirs(self.save_folder, exist_ok=True)

    def retrieve_paper(self, page=None):
        # page参数是可选的，如果传入，表示使用的是playwright获取页面内容
        # 如果不传入，则使用requests获取页面内容
        # get_paper_list方法和get_paper_info方法需要实现，具体实现由子类实现
        try:
            paper_list = self.get_paper_list(page) if page else self.get_paper_list()
            if not paper_list:
                self.logger.error("错误：没有找到任何文章，请等待下一次尝试。")
                return False

            new_paper_list = []
            for paper in paper_list:
                category, title, datetime, href = (
                    paper["category"],
                    paper["title"],
                    paper["pubtime"],
                    paper["href"],
                )

                if category == "政策解读库" and title.startswith("图解"):
                    self.logger.info(f"【跳过图解】[{category}] {datetime} {title} ...")
                    continue

                name_pure = f"{category}-{self.format_date(datetime)}-{title}"
                # 处理文件名，去除特殊字符，替换为下划线
                rstr = r"[\/\\\:\*\?\"\<\>\|]"  # '/ \ : * ? " < > |'
                name_pure = re.sub(rstr, "_", name_pure)
                # 下载正文
                file_name = f"{name_pure}.docx"
                if not self.is_downloaded(file_name):
                    self.logger.info(f"【正在处理】[{category}] {datetime} {title} ...")
                    p_elements = (
                        self.get_paper_info(href, page)
                        if page
                        else self.get_paper_info(href)
                    )

                    if p_elements:
                        self.create_docx(
                            file_name, title, category, datetime, p_elements
                        )
                        new_paper_list.append(paper)

                # 下载附件
                # 使用下划线，确保排序时.在_之前
                file_name_attachement = f"{name_pure}_相关文件.docx"
                if "政策解读库" in file_name_attachement and not self.is_downloaded(
                    file_name_attachement
                ):
                    href_attachement = paper.get("href_attachement")
                    if href_attachement:
                        self.logger.info(
                            f"【正在处理】[{category}] {datetime} {title} 相关文件 ..."
                        )
                        # 下面的函数在browser子类中定义
                        self.retrieve_attachement(
                            page,
                            href_attachement,
                            file_name_attachement,
                            category,
                            datetime,
                        )
            if self.notify_switch:
                self.notify(new_paper_list, len(paper_list))
            return True

        except Exception as e:
            self.logger.error(f"retrieve_paper()运行过程出错：{str(e)}")
            return False

    def is_downloaded(self, file_name):
        return any(file_name in files for _, _, files in os.walk(self.save_folder))

    @staticmethod
    def set_page_to_a4(doc):
        for section in doc.sections:
            section.page_height = Mm(297)
            section.page_width = Mm(210)

    def extract_and_write_title(self, doc, title, category, datetime):
        self.format_title(doc, category, "楷体", WD_ALIGN_PARAGRAPH.LEFT)
        self.format_title(doc, title, "黑体", WD_ALIGN_PARAGRAPH.CENTER, Pt(16))
        self.format_title(doc, datetime, "黑体", WD_ALIGN_PARAGRAPH.CENTER)

    @staticmethod
    def format_title(doc, text, font_name, alignment, font_size=None):
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if font_size:
            run.font.size = font_size

    @staticmethod
    def format_paragraph(doc):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 设置两端对齐
        paragraph.paragraph_format.first_line_indent = Pt(11 * 2)
        return paragraph

    @staticmethod
    def format_paragraph_font(paragraph):
        run = paragraph.add_run()
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        return run

    def create_docx(self, file_name, title, category, datetime, p_elements):
        doc = Document()  # 创建一个新的Word文档
        self.set_page_to_a4(doc)  # 设置页面大小为A4
        self.extract_and_write_title(doc, title, category, datetime)
        self.extract_and_write_paragraphs(doc, p_elements)
        doc_path = os.path.join(self.save_folder, file_name)
        doc.save(doc_path)
        self.change_file_owner(doc_path)

    def format_date(self, date_time):
        match = re.search(r"^(\d{4}-\d{2}-\d{2})", date_time)
        return match.group(1) if match else None

    def notify(self, new_paper_list, total):
        # 发送通知函数，利用requests库发送http post把payload发送到webhook地址http://192.168.1.2:1880/scraper-news
        # 这里的payload可以包含新闻的标题、链接、发布时间、分类等信息
        # 发送到nodered，处理后再发送到ha的通知

        # 如果没有新文章，则不发送通知
        # 这里的total是所有文章的数量，new_paper_list是新增的文章列表
        if (count := len(new_paper_list)) == 0:
            return

        self.logger.info(
            f"{self.source_name}新增{count}篇文章，跳过{total-count}篇已存在的文章。"
        )
        ha_notify = {
            "title": f"【{self.source_name_cn}】更新{len(new_paper_list)}条",
            "message": "\n".join(
                [f"{paper['category']}-{paper['title']}" for paper in new_paper_list]
            ),
        }
        payload = {"notify": ha_notify}
        if self.webhook_url:
            # requests不需要等待返回，用子进程直接发送

            self.logger.info(f"发送消息到webhook: {self.webhook_url}")
            requests.post(self.webhook_url, json=payload)

    def change_file_owner(self, file_path):
        """
        更改文件所有者, 仅在Linux系统下有效
        """
        try:
            chown = getattr(os, "chown", None)
            if chown:
                chown(file_path, self.uid, self.gid)
        except Exception as e:
            self.logger.info(f"更改文件所有者失败: {e}")

    # ------------------------------------------------------- #
    # 子类需要实现的抽象方法，为了防止子类未实现，这里会raise NotImplementedError
    def extract_and_write_paragraphs(self, doc, p_elements):
        # 必须在子类中实现
        raise NotImplementedError("子类必须实现extract_and_write_paragraphs方法")

    def get_paper_info(href, page=None):
        # 必须在子类中实现
        raise NotImplementedError("子类必须实现get_paper_info方法")

    def get_paper_list(self, page=None):
        # 必须在子类中实现
        raise NotImplementedError("子类必须实现get_paper_list方法")

    def retrieve_attachement(
        self, page, href_attachement, file_name_attachement, category, datetime
    ):
        # 必须在子类中实现
        raise NotImplementedError("子类必须实现retrieve_attachement方法")
