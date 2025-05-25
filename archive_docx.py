import os
from datetime import datetime
from docx import Document
import shutil
import sys
from utils_func import (
    load_config,
    mkdirs_with_owner,
    change_file_owner,
    setup_logging,
    mqtt_publish,
)
from syno_drive_orgnizer import process_stars_move_api
import time

# 调用时传入参数: move、combine或者stars


class DocxArchiver:
    def __init__(self, config_path="config/config.yaml"):
        self.logger = setup_logging()
        self.config = load_config(config_path)
        self.uid = self.config.get("uid")
        self.gid = self.config.get("gid")
        self.root_folder = os.path.join(os.getcwd(), self.config.get("save_folder"))
        self.news_sites = self.config.get("news_sites")
        self.current_date = datetime.now().strftime("%Y-%m-%d")

    def merge_docx_files(
        self, docx_folder, output_folder, file_prefix, break_flag=True, by_category=True
    ):
        # 获取输入文件夹中的所有docx文件
        docx_files = sorted([f for f in os.listdir(docx_folder) if f.endswith(".docx")])
        # 获取文件数量以避免多次计算
        num_files = len(docx_files)
        if num_files == 0:
            self.logger.info(f"【{file_prefix}】没有找到任何docx文件")
            return

        # 创建一个字典用于存储不同类别的文档内容（如果按类别合并）
        category_documents = {} if by_category else None
        merged_document = Document() if not by_category else None

        for index, docx_file in enumerate(docx_files):
            # 从文件名中提取类别
            category_current = docx_file.split("-")[0]
            docx_path = os.path.join(docx_folder, docx_file)
            document = Document(docx_path)

            if break_flag and index < num_files - 1:
                if by_category:
                    # 获取下一个文件的类别
                    category_next = docx_files[index + 1].split("-")[0]
                    # 如果当前类别与下一个文件类别相同, 则添加分页符
                    if category_current == category_next:
                        document.add_page_break()
                else:
                    document.add_page_break()
            else:
                document.add_paragraph()

            if by_category:
                # 如果该类别还没有Document对象, 创建一个新的Document对象
                if category_current not in category_documents:
                    category_documents[category_current] = Document()
                # 获取当前类别的Document对象
                merged_document = category_documents[category_current]

            # 将当前文档的内容添加到合并文档中
            for element in document.element.body:
                merged_document.element.body.append(element)

        # 保存每个类别的合并文档或单个合并文档
        mkdirs_with_owner(output_folder, self.uid, self.gid)

        if by_category:
            for category_current, merged_document in category_documents.items():
                file_name = f"{file_prefix}-{category_current}-{self.current_date}"
                file_name = (
                    f"{self._generate_unique_file_name(output_folder, file_name)}.docx"
                )
                output_file = os.path.join(output_folder, file_name)
                merged_document.save(output_file)
                change_file_owner(output_file, self.uid, self.gid)
                self.logger.info(f"合并文件 {file_name} 成功")
        else:
            file_name = f"{file_prefix}-{self.current_date}"
            file_name = (
                f"{self._generate_unique_file_name(output_folder, file_name)}.docx"
            )
            output_file = os.path.join(output_folder, file_name)
            merged_document.save(output_file)
            change_file_owner(output_file, self.uid, self.gid)
            self.logger.info(f"合并文件 {file_name} 成功")

    def _generate_unique_file_name(self, output_folder, file_name):
        output_file = os.path.join(output_folder, f"{file_name}.docx")
        if os.path.exists(output_file):
            return self._generate_unique_file_name(output_folder, f"{file_name}-new")
        return file_name

    def move_docx_files_to_archive(self, docx_folder, target_folder):
        mkdirs_with_owner(target_folder, self.uid, self.gid)
        num = 0
        for file in os.listdir(docx_folder):
            if file.endswith(".docx"):
                num += 1
                shutil.move(
                    os.path.join(docx_folder, file),
                    os.path.join(target_folder, file),
                )
        if num > 0:
            self.logger.info(f"移动 {num} 个文件到 {target_folder} 成功")

    def process_combine_mode(self, break_flag=True):
        output_folder = os.path.join(self.root_folder, "合并文档-按类别")
        for item in self.news_sites:
            docx_folder = os.path.join(self.root_folder, item.get("name"))
            file_prefix = item.get("name_cn")

            self.merge_docx_files(
                docx_folder, output_folder, file_prefix, break_flag, by_category=True
            )

            archive_folder = os.path.join(
                docx_folder, "合并过的文件", self.current_date
            )
            self.move_docx_files_to_archive(docx_folder, archive_folder)

    def process_stars_mode(self, break_flag=True):
        output_folder = os.path.join(self.root_folder, "合并文档-加星")
        for item in self.news_sites:
            docx_folder = os.path.join(self.root_folder, item.get("name"), "加星")
            if not os.path.exists(docx_folder):
                continue

            file_prefix = f"{item.get('name_cn')}-加星"
            self.merge_docx_files(
                docx_folder, output_folder, file_prefix, break_flag, by_category=False
            )

            archive_folder = os.path.join(
                self.root_folder,
                item.get("name"),
                "合并过的文件",
                f"加星-{self.current_date}",
            )
            self.move_docx_files_to_archive(docx_folder, archive_folder)

    def process_move_mode(self):
        for item in self.news_sites:
            docx_folder = os.path.join(self.root_folder, item.get("name"))
            archive_folder = os.path.join(docx_folder, "已读的文件", self.current_date)
            self.move_docx_files_to_archive(docx_folder, archive_folder)

    def run(self, mode, break_flag=True):
        if mode not in ["move", "combine", "stars"]:
            raise ValueError("参数错误, 请传入参数: move、combine 或者 stars")

        # 移动加星文件到指定文件夹, 防止掉加星
        self.logger.info("开始移动加星文件到指定文件夹")
        process_stars_move_api()

        if mode == "combine":
            self.process_combine_mode(break_flag)
        elif mode == "stars":
            self.process_stars_mode(break_flag)
        elif mode == "move":
            self.process_move_mode()


if __name__ == "__main__":
    # 示例用法
    args = sys.argv
    if len(args) < 2 or len(args) > 3:
        print("参数错误, 请传入参数: move、combine 或者 stars")
        sys.exit(1)

    mode = args[1]
    break_flag = True  # 默认情况下设置为 True, 即合并文档时添加分页符
    if len(args) == 3 and args[2] == "nobreak":
        break_flag = False

    try:
        archiver = DocxArchiver()
        archiver.run(mode, break_flag)
        time.sleep(2)
        mqtt_publish(payload="update-drive")
    except Exception as e:
        print(f"发生错误: {str(e)}")
        sys.exit(1)
